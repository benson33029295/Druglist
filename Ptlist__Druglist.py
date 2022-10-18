# -*- coding: utf-8 -*-
"""
Created on Tue Apr 26 16:35:24 2022

@author: mrd
"""


print("Author: M117 H.R.T.")
print("Revised on 2022-04-26")
print("#"*40)

doc_num = input("DOC帳號:")
password_num = input("密碼:")
print("#"*20)
#name = input("名稱:")


section_num = input("科別:")
vs_num = input("主治DOC:")
ward_num = input("病房:")
bed_num = ""

chart_num = ""
chinaname = ""

'''
doc_num = ""
password_num = ""
section_num = ""
vs_num = "DOC10559"
ward_num = ""
bed_num = ""
chart_num = "3352700"
chinaname = ""
'''
print("#"*40)



def tableformat(tablename, fontname, fontsize):
    for row in tablename.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.bold = False #加粗
                    run.font.italic = False #傾斜 等等...
                    run.font.name = fontname
                    run.font.size = Pt(fontsize)
                    run.font.color.rgb = RGBColor(0, 0, 0)

def tablecellformat(cell, fontname, fontsize, bold=True):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.bold = bold #加粗
            run.font.italic = False #傾斜 等等...
            run.font.name = fontname
            r = run._element.rPr.rFonts
            r.set(qn("w:eastAsia"), fontname)
            run.font.size = Pt(fontsize)
            run.font.color.rgb = RGBColor(0, 0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    
def set_table_different_width(table, widths):
    for x, width in enumerate(widths):
        for cell in table.columns[x].cells:
            cell.width = Cm(width)


#https://blog.csdn.net/lly1122334/article/details/109669667
def add_paragraph(document_or_cell, content="", font_name="Calibri", font_size=6,
                  bold=False, italic=False, underline=False, strike=False,
                  font_color=None, highlight_color=None):
    paragraph = document_or_cell.add_paragraph()
    run = paragraph.add_run(str(content))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.strike = strike #字體刪除線
    
    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)
    if highlight_color:
        run.font.highlight_color = highlight_color
        
def add_run(paragraph, content="", font_name="Calibri", font_size=6,
                  bold=False, italic=False, underline=False, strike=False,
                  font_color=None, highlight_color=None):
    #paragraph = document_or_cell.add_paragraph()
    run = paragraph.add_run(str(content))
    run.font.name = font_name
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.strike = strike #字體刪除線
    
    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)
    if highlight_color:
        run.font.highlight_color = highlight_color
'''
WD_COLOR_INDEX.YELLOW
AUTO = 'default'
BLACK = 'black'
BLUE = 'blue'
BRIGHTGREEN = 'green'
DARKBLUE = 'darkBlue'
DARKRED = 'darkRed'
DARKYELLOW = 'darkYellow'
GRAY25 = 'lightGray'
GRAY50 = 'darkGray'
GREEN = 'darkGreen'
PINK = 'magenta'
RED = 'red'
TEAL = 'darkCyan'
TURQUOISE = 'cyan'
VOILET = 'darkMagenta'
WHITE = 'white'
YELLOW = 'yellow'
'''
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def delete_paragraph(paragraph):
    p = paragraph._element
    tc = p.getparent()
    tc.remove(p)
    p._p = None
    p._element = None

def del_blankstring_paragraph(document_or_cell):
    for paragraph in document_or_cell.paragraphs:
        #print(paragraph.text)
        if paragraph.text == "" :
            delete_paragraph(paragraph) 


def aligns_last(string, length = 73):
    difference = length - len(string) 
    
    new_string = ""
    space = " "
    # 計算限定長度為20時需要補齊多少個空格
    if difference == 0: # 若差值為0則不需要補
        return string
    
    
    elif difference < 0: 
        print('錯誤：限定的對齊長度小於字元串長度!') 
        return None
    
    
    else: 
        new_string = string[:-31] + space * difference + string[-31:]
        return new_string
    # 若是全形，則不轉換 return new_string + space*(difference) 
    # 返回補齊空格後的字元串 





def aligns_long(string, length = 86):
    difference = length - len(string) 
    
    new_string = ""
    space = " "
    # 計算限定長度為20時需要補齊多少個空格
    if difference == 0: # 若差值為0則不需要補
        return string
    
    
    elif difference < 0: 
        print('錯誤：限定的對齊長度小於字元串長度!') 
        return None
        
    
    
    else: 
        new_string = string[:-44] + space * difference + string[-44:]
        return new_string

def aligns_drugname(string, length = 20):
    difference = length - len(string) 
    
    new_string = ""
    space = " "
    # 計算限定長度為20時需要補齊多少個空格
    if difference == 0: # 若差值為0則不需要補
        return string
    
    
    elif difference < 0: 
        print('錯誤：限定的對齊長度小於字元串長度!') 
        return string[:20]
    
    
    else: 
        new_string = string + space * difference
        return new_string
    # 若是全形，則不轉換 return new_string + space*(difference) 
    # 返回補齊空格後的字元串 
    
    
    

import bs4

import requests
import pandas as pd
#import os
import re
import math
##########################################################################
import random
import time
printtime = time.strftime("%Y-%m-%d(%A) %H:%M ", time.localtime())
 
delay_choices = [0.5, 1]  #延遲的秒數
delay = random.choice(delay_choices)  #隨機選取秒數
#time.sleep(delay)  #延遲

##########################################################################


def validateTitle(title):
    rstr = r"[\/\\\:\*\?\"\<\>\|.]"  # '/ \ : * ? " < > |'
    new_title = re.sub(rstr, "", title)  # 替換為...下劃線
    return new_title

##########################################################################


referlink_choices = ["www.tsgh.ndmctsgh.edu.tw","www.ndmctsgh.edu.tw","www.medicaltravel.org.tw","www.google.com.tw"]
referlink = random.choice(referlink_choices)

##########################################################################
'''
from fake_useragent import UserAgent

user_agent = UserAgent()
'''
#response = requests.get(url="https://example.com", headers={ 'user-agent': user_agent.random })
##########################################################################

login_session = requests.Session()
# 按照網頁的請求構造請求頭
headers = {
        #'User-Agent': user_agent.random,
        'Connection': 'keep-alive',
        #'Referer': referlink,
        'Origin': 'http://f5-ws5.ndmctsgh.edu.tw',
        'Referer': 'http://f5-ws5.ndmctsgh.edu.tw/eForm/Account/Login'
    }



data = {
    'login_id': doc_num,
    'password': password_num
}
# 按照服務端的需要構造請求，並獲得session，記錄在login_session
response = login_session.post('http://f5-ws5.ndmctsgh.edu.tw/eForm/Account/Login', data=data, verify=False, headers=headers)
#print(response.content)





# 登入後，我們需要獲取另一個網頁中的內容
#response2 = login_session.get('https://www-nature-com.autorpa.ndmctsgh.edu.tw/articles/s41580-021-00417-y',headers = headers)
#print(response2.status_code)
#print(response2.text)





##########################################################################
print("aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa")

################################################################




ptsearchdata = {
        'SearchChartno': chart_num,
        'SearchChinaname': chinaname,
        'SearchIDNo': '', #身分證號
        'SearchSectionNo': section_num,
        'SearchNRCode': ward_num,
        'SearchBedCode': bed_num,
        'SearchVSDR': vs_num,
        'SearchSpecialNote': '',
        'SearchCcGroup[]': ''
}

htmlfile1 = login_session.post('http://f5-ws5.ndmctsgh.edu.tw/eForm/Patient/Result', data=ptsearchdata, verify=False, headers=headers)
#print(response.content)
 

if htmlfile1.status_code == requests.codes.ok:
    print("取得成功")
else:
    print("取得失敗")
    
print(htmlfile1.text)

pt_list = htmlfile1.json()

print(type(pt_list))


outdates=[]
chartnums=[]
#六碼者labdataurl會有Q   ===> ok
hcasenums=[]

NameGenderAges = []
NrBedNos = []
indatetimes = []
vsnames = []


for x in pt_list:
    outdate1 = x['OUTDATETIME']
    print("出院日" + outdate1)
    outdates.append(outdate1)
    
    if outdate1 == "":
        chartnum1 = x['CHARTNO']
        print("病歷號" + chartnum1)
        chartnums.append(chartnum1)
        
        hcasenum1 = x['HCASENO']
        print(hcasenum1)
        hcasenums.append(hcasenum1)
        
        NameGenderAge1 = x['NameGenderAge']
        NrBedNo1 = x['NrBedNo']
        indatetime1 = x['INDATETIME'][:7]
        vsname1 = x['VSDRNAME']
        NameGenderAges.append(NameGenderAge1)
        NrBedNos.append(NrBedNo1)
        indatetimes.append(indatetime1)
        vsnames.append(vsname1)



         

patients_drugs=[]

patients_drugs_forsearch = []
for h in range(len(hcasenums)):
    print("drug" + str(h)*10)     
    ###################################################################################################
    #drug
    
    drug_url = 'http://mobilereport.ndmctsgh.edu.tw/mr/HISEXNDREPORT.aspx?login_id=' + doc_num + '&special=n&cno=' + chartnums[h]
    
    htmlfile6 = login_session.get(drug_url, verify=False, headers=headers)
    #print(response.content)
     
    
    if htmlfile6.status_code == requests.codes.ok:
        print("取得成功")
    else:
        print("取得失敗")
        
    #print(htmlfile4.text)
    
    #import bs4
    objSoup6=bs4.BeautifulSoup(htmlfile6.text, 'lxml')
    print(type(objSoup6))
    
    
    #div data-role="content"
    drugdata = objSoup6.find('div', {'data-role': 'content'})
    #print("資料型態", type(table))
    #print("串列長度", len(table))
    
    print(drugdata.text[87:-1])
    
    #patients_drugs.append(drugdata.text[86:])
    
    
    
    
    druglist = drugdata.text[87:-1].splitlines()
    




    


    
    
    
    
    
    drugs= []
    
    drug_dose = []
    
    drug_freq = []
    
    drug_method = []
    
    drug_start_time = []
    #drugothers = []
    
    #drugs_forsearch = []
    
    
    for drugline in druglist:
        #diff = 84-len(drugline)
        
        if len(drugline) > 75:
            drugline = aligns_long(drugline)
        else:
            drugline = aligns_last(drugline)
        
        
        
        
        

        '''
        #drugname = drugline[:30]
        drugname_forsearch = drugline[:42]
        drugs_forsearch.append(drugname_forsearch)
        '''
        '''
        drugother = drugline[43:62]  #686
        drugothers.append(drugother)
        '''
        drugname = drugline[:42].title()
        drugs.append(drugname)
        
        #686
        drug_dose1 = drugline[42:48]
        drug_dose.append(drug_dose1)
        drug_freq1 = drugline[48:56]
        drug_freq.append(drug_freq1)
        drug_method1 = drugline[56:62]
        drug_method.append(drug_method1)
        
        drug_start_time1 = drugline[62:73]
        drug_start_time.append(drug_start_time1)
        #drugs.append(drugname + drugother)
        
        
    #drugdata_cut = '\n'.join(drugs)
    #patients_drugs.append(drugdata_cut)
    
    druginfo_dataframe = pd.DataFrame(list(zip(drugs, drug_dose, drug_freq, drug_method, drug_start_time)),\
                                 columns=["商品名", "dose", "freq", "method", "開始時間"])
    
    
    
    
    
    
    
    
    patients_drugs.append(druginfo_dataframe)
    
    patients_drugs_forsearch.append(drugs)

    






##############################################################################

#df =  pd.read_csv("drug.csv", usecols=["drug"], encoding="utf-8")




#druglist = df["drug"].tolist()


drugtypelist=['Alimentary tract and metabolism',
'Blood and blood forming organs',
'Cardiovascular system',
'Dermatologicals',
'Genito urinary system and sex hormones',
'Systemic hormonal preparations, excl. sex hormones and insulins',
'Antiinfectives for systemic use',
'Antineoplastic and immunomodulating agents',
'Musculo-skeletal system',
'Nervous system',
'Antiparasitic products, insecticides and repellents',
'Respiratory system',
'Sensory organs',
'Various',
'None!'
]

drugindexlist=[4.0,
9.0,
2.0,
8.0,
5.0,
6.0,
10.2,
10.4,
8.0,
1.0,
10.2,
3.0,
8.0,
100.0,
1000.0
]

drugindex_color_list=['92D050',
'00B050',
'FFC000',
'C6D412',
'00B0F0',
'FF6699',
'9BC2E6',
'0070C0',
'C6D412',
'FF0000',
'9BC2E6',
'FFFF00',
'C6D412',
'7030A0',
'A6A6A6'
]

h_dataframes = []

for h in range(len(hcasenums)):
    print("drug" + str(h)*10)   
    
    
   


    objtagdrugnames = []
    objtagdrugnames_short = []

    objtagdrugtype1s = []
    drugindex = []
    objtagdrugindis = []
    
    
    for i in range(len(patients_drugs_forsearch[h])):
        print("drug" + str(h)*10 + str(i)*10)  

        
        
        #querydrug = druglist[i][:30]
        querydrug = patients_drugs_forsearch[h][i]
        
        
        print(querydrug)
        url = "http://eserver101.ndmctsgh.edu.tw/Med/web/MedList.aspx?QryType=1&PrefixName=" + querydrug
        
        
        #import requests
        #import bs4
        
        
        #抓藥物網址
        #time.sleep(delay)
        response = requests.get(url)
        #print(response.text)
        
        objSoup1=bs4.BeautifulSoup(response.text, 'lxml')
        print(type(objSoup1))
        
        
        objtag1 = objSoup1.find("td")
        
        
        if objtag1 == None:
            objtagdrugnames.append("None!")
            objtagdrugnames_short.append("None!")

            objtagdrugtype1s.append("None!")
            #h_drugindex[h].append("None!")
            drugindex.append(1000.0)
            objtagdrugindis.append("None!")
            continue
            
        
        objtag2 = objtag1.find("a")
        
        drugurl = "http://eserver101.ndmctsgh.edu.tw/Med/web/" + objtag2.get("href")
        
        print(drugurl)
        print("a" *20)
        
        #抓藥物資料
        
        response2 = requests.get(drugurl)
        print(response.text)
        
        objSoup2=bs4.BeautifulSoup(response2.text, 'lxml')
        print(type(objSoup2))
        
        print("b" *20)
        
        objtag3 = objSoup2.find('div', {'id': 'accordion'})
        
        
        '''
        objtag3 = objSoup2.find('div', {'id': 'accordion'}).children
        print(type(objtag3))
        for child in objtag3:
            #print(child.text)
            print("分隔"*5)
        '''
        
        
        
        
        #學名
        print("學名" *10)
        # <span id="labDrugElemCode1">
        objtagdrugname1 = objtag3.find('span', {'id': 'labDrugElemCode1'})
        objtagdrugname2 = objtag3.find('span', {'id': 'labDrugElemCode2'})
        objtagdrugname3 = objtag3.find('span', {'id': 'labDrugElemCode3'})
        objtagdrugname4 = objtag3.find('span', {'id': 'labDrugElemCode4'})
        #print(type(objtag5))
        #print(len(objtag5))
        '''
        if objtagdrugname1 == None:
            objtagdrugnames.append("None!")
            objtagdrugtype1s.append("None!")
            #h_drugindex[h].append("None!")
            drugindex.append(1000.0)
            objtagdrugindis.append("None!")
            continue
        '''

        
        objtagdrugname_1234 = objtagdrugname1.text +" "+ objtagdrugname2.text +" "+ objtagdrugname3.text +" "+ objtagdrugname4.text
        print(objtagdrugname_1234.strip())
        
        if objtagdrugname_1234.strip() == "":
            objtagdrugnames.append("None!")
            objtagdrugnames_short.append("None!")

            objtagdrugtype1s.append("None!")
            #h_drugindex[h].append("None!")
            drugindex.append(1000.0)
            objtagdrugindis.append("None!")
            continue
        
        objtagdrugnames.append(objtagdrugname_1234.strip().title())
        objtagdrugnames_short.append(aligns_drugname(objtagdrugname_1234.strip().title()))


        
        
        
        
        #分類
        print("分類" *10)
        #<a id="labAHFSA1"
        objtagdrugtype1 = objtag3.find('a', {'id': 'labATC1'})
        objtagdrugtype2 = objtag3.find('a', {'id': 'labATC2'})
        objtagdrugtype3 = objtag3.find('a', {'id': 'labATC3'})
        objtagdrugtype4 = objtag3.find('a', {'id': 'labATC4'})
        objtagdrugtype5 = objtag3.find('a', {'id': 'labATC5'})
        '''
        if objtagdrugtype1 == None:
            #objtagdrugnames.append("None!")
            objtagdrugtype1s.append("None!")
            #h_drugindex[h].append("None!")
            drugindex.append(1000.0)
            objtagdrugindis.append("None!")
            continue
        '''
        
        print(objtagdrugtype1.text[2:].strip())
        
        if objtagdrugtype1.text[2:].strip() == "":
            #objtagdrugnames.append("None!")
            objtagdrugtype1s.append("None!")
            #h_drugindex[h].append("None!")
            drugindex.append(1000.0)
            objtagdrugindis.append("None!")
            continue

        
        objtagdrugtype1s.append(objtagdrugtype1.text[2:].strip())
        
        objtagdrugtype222 = objtagdrugtype2.text[4:].strip()
        
        if objtagdrugtype2.text[5:].strip() == "Drugs used in diabetes":
            drugindex.append(6.0)
        else:
            for j in range(len(drugtypelist)):
                if objtagdrugtype1.text[2:].strip()==drugtypelist[j]:
                    drugindex.append(drugindexlist[j])
        
        

        '''
        for i in range(objtag5):
            print(str(i)*10)
            print(objtag5[i].text)
        #print(objtag3.text)
        '''
        
        


        
        
        #Indication
        print("Indication" *10)
        #<span id="labDOHINDICATION"><br>高血壓。</span>
        objtagdrugindi = objtag3.find('span', {'id': 'labDOHINDICATION'})
        
        print(objtagdrugindi.text[:20])
        objtagdrugindis.append(objtagdrugindi.text)
        

        
    #重排
    patients_drugs[h]["Drugtypeindex"] = drugindex
    patients_drugs[h]["Drugtypes"] = objtagdrugtype1s
    patients_drugs[h]["Drugnames"] = objtagdrugnames
    patients_drugs[h]["Drugnames_short"] = objtagdrugnames_short

    patients_drugs[h]["Indications"] = objtagdrugindis
    
    #drugdataframe = pd.DataFrame(list(zip(patients_drugs[h], drugindex, objtagdrugtype1s, objtagdrugnames, objtagdrugindis)),\
    #                             columns=["商品名", "Drugtypeindex", "Drugtypes", "Drugnames", "Indications"])
    
    drugdataframe2 = patients_drugs[h].sort_values(by=['Drugtypeindex', "Drugnames"], ascending=True)
    
    #localtimes = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()).replace(":","")
    #drugdataframe2.to_excel('try drug' + chartnums[h] + NameGenderAges[h] + '.xls')
    
    drugdataframe2.index = drugdataframe2["Drugtypeindex"]
    #排對再append
    
    drugdataframe3 = drugdataframe2[["Drugtypeindex", "商品名", "Drugnames_short", "dose", "freq", "method", "開始時間"]]
    h_dataframes.append(drugdataframe3)








#drugdataframe2.to_excel('try drug' + localtimes + '.xls')






    
    
print("ok")
###################################################################################################

'''
htmlfile2 = login_session.get('http://f5-ws5.ndmctsgh.edu.tw/eForm/PL', verify=False, headers=headers)
#print(response.content)
 

if htmlfile2.status_code == requests.codes.ok:
    print("取得成功")
else:
    print("取得失敗")
    
#print(htmlfile2.text)

'''







import docx

from docx.shared import RGBColor
from docx.shared import Cm, Pt#, Inches  #加入可調整的 word 單位
#from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
#from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT#, WD_TABLE_ALIGNMENT

#from docx.oxml import parse_xml
from docx.oxml.ns import qn#, nsdecls



document = docx.Document()
document.styles["Normal"].font.name = "Calibri"
document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u'標楷體')




section = document.sections[0]


#調整文件左右上下邊界至 1 cm
section.left_margin=Cm(1)
section.right_margin=Cm(1)
section.top_margin=Cm(0.5)
section.bottom_margin=Cm(0.5)

#import time
#savetime = time.ctime()

#xtime = time.localtime()
#printtime = xtime[0] + xtime[1] + xtime[2]




'''
# Choosing the top most section of the page
#section = document.sections[0]
 
# Calling the header
header = section.header
 
# Calling the paragraph already present in
# the header section
header_para = header.paragraphs[0]
 
# Adding text in the header
#header_para.text = "\tPatientList:  " + printtime + "\t"
 
document.sections[0].different_first_page_header_footer = False

footer = section.footer
footer_para = footer.paragraphs[0]

'''


tablenum = math.ceil(len(hcasenums) / 5)
modenum = len(hcasenums) % 5

#for h in range(len(hcasenums)): 
    
    #print("word" + str(h)*10)
    
    #document.add_heading(NrBedNos[h] + ":  " + NameGenderAges[h] + ",  " + chartnums[h], 1) #直接添加標題
    #document.add_paragraph("入院日期: " + indatetimes[h] + ",  主治醫師: " + vsnames[h])
    
'''  

def write_ptdata():
    table.allow_autofit = True
    
    #h = g * 5 + i
    p = table.cell(2*i, 0).add_paragraph(NrBedNos[h] + "\n" + NameGenderAges[h] + "\n" + chartnums[h])# + "\n入院日期: " + indatetimes[h] + "\n主治醫師: " + vsnames[h]
    tablecellformat(table.cell(2*i, 0), "標楷體", 14, bold=True)
    
    add_run(p, content="\n入院日期: " + indatetimes[h][3:] + ", 主治醫師:"  + vsnames[h], font_name="標楷體", font_size=4,
                      bold=False, italic=False, underline=False, strike=False,
                      font_color=None, highlight_color=None)
    table.cell(2*i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  #垂直對齊，置中
    
    
    cellone_two = "Temp:_____-_____\nPulse:_____-_____\nRespRate:_____-_____\nBlood Pressure:_____/_____-_____/_____\nSpO2:_____-_____\nInput:______\n\nOutput:______"
    table.cell(2*i, 1).add_paragraph(cellone_two)
    tablecellformat(table.cell(2*i, 1), "Calibri", 6, bold=True)
    
    table.cell(2*i, 2).add_paragraph("Lab:\n\n\n\n\nImage:")
    tablecellformat(table.cell(2*i, 2), "Calibri", 6, bold=True)
    
    table.cell(2*i, 3).add_paragraph("Drug:")
    tablecellformat(table.cell(2*i, 3), "Calibri", 6, bold=True)
    
    table.cell(1+2*i, 1).add_paragraph("CC:" + " "*48 + "Signs: \n\n吃喝: \n\n尿: \n\n便: \n\n睡: ")
    tablecellformat(table.cell(1+2*i, 1), "Calibri", 6, bold=True)
    

    
    #for cell in table.row_cells(2*i):  #第一列cell
    #    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  #垂直對齊，置中
    for j in range(4):
        del_blankstring_paragraph(table.cell(2*i, j))
        del_blankstring_paragraph(table.cell(1+2*i, j))




for g in range(tablenum):
    
    #header_para.text = "\tPatientList:  " + printtime + "\t"
    #footer_para.text = "\t" + str(g+1) + "\t"
     
    
    document.add_paragraph("PatientList: " + printtime + "          Page " + str(g+1)) 
    table = document.add_table(10, 4)
    table.style = "Table Grid"
    for row in table.rows:
        row.height = Cm(2.5)
    
    
    if g < (tablenum-1):
        
        for i in range(5):
            write_ptdata()

    elif g == (tablenum-1):  
        if modenum == 0:
            for i in range(5):
                write_ptdata()
        else:
            for i in range(modenum):
                write_ptdata()
    
    document.add_page_break()
    
del_blankstring_paragraph(document)

'''



'''


for h in range(len(hcasenums)): 
    
    #header_para.text = "\tPatientList:  " + printtime + "\t"
    #footer_para.text = "\t" + str(g+1) + "\t"
     
    
    document.add_paragraph("PatientList: " + printtime)
    #document.add_paragraph(NrBedNos[h] + " "*20 + NameGenderAges[h] + " "*20 + chartnums[h])
    
    table = document.add_table(10, 4)
    table.style = "Table Grid"
    for row in table.rows:
        row.height = Cm(2.5)
    
    
    p = table.cell(0, 0).add_paragraph(NrBedNos[h] + "\n" + NameGenderAges[h] + "\n" + chartnums[h])# + "\n入院日期: " + indatetimes[h] + "\n主治醫師: " + vsnames[h]
    tablecellformat(table.cell(0, 0), "標楷體", 14, bold=True)
    
    
    add_run(p, content="\n入院日期: " + indatetimes[h][3:] + ", 主治醫師:"  + vsnames[h], font_name="標楷體", font_size=4,
                      bold=False, italic=False, underline=False, strike=False,
                      font_color=None, highlight_color=None)
    table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  #垂直對齊，置中
    del_blankstring_paragraph(table.cell(0, 0))
    
    
    for i in range(5):
        
        cellone_two = "Temp:_____-_____\nPulse:_____-_____\nRespRate:_____-_____\nBlood Pressure:_____/_____-_____/_____\nSpO2:_____-_____\nInput:______\n\nOutput:______"
        table.cell(2*i, 1).add_paragraph(cellone_two)
        tablecellformat(table.cell(2*i, 1), "Calibri", 6, bold=True)
        
        table.cell(2*i, 2).add_paragraph("Lab:\n\n\n\n\nImage:")
        tablecellformat(table.cell(2*i, 2), "Calibri", 6, bold=True)
        
        table.cell(2*i, 3).add_paragraph("Drug:")
        tablecellformat(table.cell(2*i, 3), "Calibri", 6, bold=True)
        
        table.cell(1+2*i, 1).add_paragraph("CC:" + " "*20 + "Signs: \n\n吃喝: \n\n尿: \n\n便: \n\n睡: ")
        tablecellformat(table.cell(1+2*i, 1), "標楷體", 6, bold=True)
        #tablecellformat(table.cell(1+2*i, 1), "Calibri", 6, bold=True)
    
        del_blankstring_paragraph(table.cell(2*i, 1))
        del_blankstring_paragraph(table.cell(2*i, 2))
        del_blankstring_paragraph(table.cell(2*i, 3))
        del_blankstring_paragraph(table.cell(1+2*i, 1))
    
    #for row in table.rows:
    #    for cell in row.cells:
    #        del_blankstring_paragraph(cell)
    
    document.add_page_break()
    
del_blankstring_paragraph(document)

'''

print("分隔線分隔線分隔線分隔線分隔線分隔線分隔線分隔線分隔線")

for h in range(len(hcasenums)): 
    
    if h % 20 == 0:
        document.add_paragraph("PatientList: " + printtime) 
        table = document.add_table(5, 4)
        table.style = "Table Grid"
        for row in table.rows:
            row.height = Cm(5)
        

    ptinfo = NrBedNos[h] + "______" + NameGenderAges[h] + "______" + chartnums[h] + "\n" #+ "\n入院日期: " + indatetimes[h] + "     主治醫師: " + vsnames[h]
    

    #####ij定義
    modenum = h % 20
    j = math.floor(modenum / 5) #= columnnum
    i = modenum % 5  #= rownum 

    
    
    p = table.cell(i, j).add_paragraph(NrBedNos[h] + "\n" + NameGenderAges[h] + "\n" + chartnums[h])# + "\n入院日期: " + indatetimes[h] + "\n主治醫師: " + vsnames[h]
    tablecellformat(table.cell(i, j), "標楷體", 18, bold=True)
    
    
    add_run(p, content="\n入院日期: " + indatetimes[h][3:] + ", 主治醫師:"  + vsnames[h], font_name="標楷體", font_size=8,
                      bold=False, italic=False, underline=False, strike=False,
                      font_color=None, highlight_color=None)
    table.cell(i, j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  #垂直對齊，置中
    
    del_blankstring_paragraph(table.cell(i, j))

    if h % 20 == 19:
        document.add_page_break()


document.add_page_break()


print("分隔線分隔線分隔線分隔線分隔線分隔線分隔線分隔線分隔線")



for h in range(len(hcasenums)): 
    if h % 10 == 0:
        document.add_paragraph("DrugList: " + printtime) 
        table = document.add_table(5, 2)
        table.style = "Table Grid"
        for row in table.rows:
            row.height = Cm(5)
        
    
    casedruglist = h_dataframes[h].values.tolist()
    
    druglines = []
    for infolist in casedruglist:
        drugline1 = " ".join(str(i) for i in infolist)
        druglines.append(drugline1)
    
    druglines_forprint = "\n".join(druglines)
    
    
    
    #document.add_paragraph(NrBedNos[h] + "______" + NameGenderAges[h] + "______" + chartnums[h] + "\n入院日期: " + indatetimes[h] + "     主治醫師: " + vsnames[h])
    #document.add_paragraph(druglines_forprint)
    ptinfo = NrBedNos[h] + "______" + NameGenderAges[h] + "______" + chartnums[h] + "\n" #+ "\n入院日期: " + indatetimes[h] + "     主治醫師: " + vsnames[h]
    
    
    
    
    #####ij定義
    modenum = h % 10
    j = math.floor(modenum / 5) #= columnnum
    i = modenum % 5  #= rownum 

    
    
    p = table.cell(i, j).add_paragraph(ptinfo)# + "\n入院日期: " + indatetimes[h] + "\n主治醫師: " + vsnames[h]
    tablecellformat(table.cell(i, j), "標楷體", 10, bold=True)
    
    
    add_run(p, content=druglines_forprint, font_name="Consolas", font_size=4,
                      bold=False, italic=False, underline=True, strike=False,
                      font_color=None, highlight_color=None)
    #table.cell(2*i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  #垂直對齊，置中
    
    
    #table.cell(i, j).add_paragraph(druglines_forprint)
    #tablecellformat(table.cell(i, j), "Consolas", 4, bold=False)
    
    
    del_blankstring_paragraph(table.cell(i, j))
    
    
    '''
    add_paragraph(document, content=ptinfo, font_name="標楷體", font_size=10,
                      bold=True, italic=False, underline=False, strike=False,
                      font_color=None, highlight_color=None)
    
    
    
    add_paragraph(document, content=druglines_forprint, font_name="Consolas", font_size=4,
                      bold=False, italic=False, underline=False, strike=False,
                      font_color=None, highlight_color=None)

    '''

    if h % 10 == 9:
        document.add_page_break()




document.add_page_break()







print('段落數量： ', len(document.paragraphs))

#document.save("try ptlist" + validateTitle(savetime) + ".docx")  # 保存文檔

#document.save("try ptlist" + printtime + ".docx")  # 保存文檔
document.save("try blank ptlist" + validateTitle(printtime) + ".docx")  # 保存文檔






